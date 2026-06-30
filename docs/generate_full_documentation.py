"""
Generate ONE full project documentation Word file.
Written in plain language so non-programmers can understand.
Run: python docs/generate_full_documentation.py
Output: docs/UgaJapa-Full-Documentation.docx
"""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
SCREENSHOTS_DIR = os.path.join(DOCS_DIR, "screenshots")
OUTPUT = os.path.join(DOCS_DIR, "UgaJapa-Full-Documentation.docx")


def setup(doc: Document) -> None:
    s = doc.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(11)
    s.paragraph_format.space_after = Pt(6)
    s.paragraph_format.line_spacing = 1.2


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def tbl(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for r in t.rows[0].cells[i].paragraphs[0].runs:
            r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()


def note(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(f"{label} ")
    r.bold = True
    p.add_run(text)


def screenshot(doc: Document, filename: str, caption: str, width_inches: float = 5.9) -> None:
    path = os.path.join(SCREENSHOTS_DIR, filename)
    if not os.path.isfile(path):
        note(doc, "Screenshot not found:", filename)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width_inches))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
    doc.add_paragraph()


def build_screenshots_section(doc: Document) -> None:
    heading(doc, "14. Screenshots — see the project working (with explanations)")
    para(doc,
         "This section shows real screenshots from the working system. You do not need to "
         "understand programming to follow them. Read each picture, then read the explanation "
         "under it.")
    para(doc,
         "The example below uses two test users in the same channel (Town Square): "
         "abdel (English) and yohana (Japanese). They are discussing a project design.")

    heading(doc, "Screenshot 1 — English user view (abdel)", 2)
    para(doc,
         "This is what abdel sees. His own messages appear in green bubbles on the right. "
         "Messages from yohana appear in white bubbles on the left. The conversation is in "
         "plain English because abdel chose English as his reading language.")
    bullets(doc, [
        "Green on the right = messages I sent.",
        "White on the left = messages someone else sent.",
        "Small speaker icon under each message = click to hear the message read aloud.",
        "The chat looks like WhatsApp, but translation is happening in the background.",
    ])
    screenshot(doc, "01-abdel-english-chat.png",
               "Figure 1 — abdel sees the chat in English (green = his messages, white = yohana's).")

    heading(doc, "Screenshot 2 — Japanese user view (yohana) — same conversation", 2)
    para(doc,
         "This is the most important screenshot for proving translation works. "
         "It is the SAME chat at the SAME time, but logged in as yohana. "
         "She sees the whole conversation in Japanese — not because yohana typed everything "
         "in Japanese, but because the system translated messages into her language automatically.")
    bullets(doc, [
        "abdel wrote in English on his screen (Screenshot 1).",
        "yohana sees abdel's words in Japanese on her screen (Screenshot 2).",
        "When yohana replies in Japanese, abdel will see her reply in English on his screen.",
        "Each person types naturally; each person reads in their own language.",
    ])
    screenshot(doc, "02-yohana-japanese-chat.png",
               "Figure 2 — yohana sees the same chat in Japanese (automatic translation).")

    heading(doc, "Screenshot 3 — Translation languages panel (how each person chooses their language)", 2)
    para(doc,
         "This panel opens from the translate icon in the channel header (top right of the chat). "
         "It is the control centre for languages. Anyone can understand it without code:")
    bullets(doc, [
        "Your receive language — the language you want incoming messages translated into.",
        "Read-aloud voice — male, female, or neutral when you click the speaker icon.",
        "Channel members — shows each person's chosen language (EN = English, JA = Japanese).",
        "The note at the bottom confirms read-aloud uses Google Text-to-Speech.",
    ])
    para(doc,
         "In this screenshot, abdel's receive language is English and yohana's badge shows JA (Japanese). "
         "That is why Screenshot 1 is English and Screenshot 2 is Japanese — same messages, "
         "different personal settings.")
    screenshot(doc, "03-translation-languages-panel.png",
               "Figure 3 — Language settings: abdel reads English, yohana reads Japanese.")

    heading(doc, "Screenshot 4 — Voice and video message buttons", 2)
    para(doc,
         "At the bottom of the chat, next to the normal message box, we added two extra buttons "
         "(shown with red arrows in the screenshot):")
    bullets(doc, [
        "Microphone — record a voice message (up to 5 minutes). The reader can click Translate to text to see the words.",
        "Video camera — record a short video message with the same idea.",
        "These work together with translation: voice can be turned into text, then translated like a normal message.",
    ])
    para(doc,
         "This shows the project is not only text translation — it also supports voice and video "
         "in a modern chat style.")
    screenshot(doc, "04-voice-video-buttons.png",
               "Figure 4 — Microphone and video buttons added to the message composer.")

    heading(doc, "What to say when showing these four screenshots to your company", 2)
    numbered(doc, [
        "Open Screenshot 1 and 2 side by side: same chat, two languages — that is the core feature.",
        "Open Screenshot 3: explain each user picks their own language once; translation is automatic after that.",
        "Open Screenshot 4: mention voice, video, and read-aloud as extra communication tools.",
        "Optional: open http://localhost:5000/health on the Translation API to confirm Google is active (Section 11).",
    ])
    note(doc, "Optional extra screenshots you can add later:",
         "API health page, or a round-trip example (English → Japanese → English) if you want even more proof.")


def build_explain_anyone_section(doc: Document) -> None:
    heading(doc, "17. How to explain this whole project to anyone (30-second version)")
    para(doc,
         "After reading this document, you should be able to explain the project without "
         "mentioning code. Here is a simple script you can use:")
    para(doc,
         '"We added translation to Mattermost, a team chat app like Slack. Each person picks '
         "the language they want to read. When someone sends a message, our plugin sends the "
         "text to our Translation API, which uses Google Translate. Each reader sees the message "
         "in their own language, but the sender still sees what they typed. We also added "
         "WhatsApp-style bubbles, a language settings panel, read-aloud, voice messages, and "
         "video messages. We did not change Mattermost itself — only installed our plugin and "
         'our separate API server."')
    heading(doc, "Slightly longer version (2 minutes)", 2)
    numbered(doc, [
        "Problem: international teams struggle when everyone speaks a different language in one chat.",
        "Solution: automatic per-user translation inside Mattermost.",
        "How: plugin in the chat + Translation API on port 5000 + Google Cloud services.",
        "Proof: Screenshots 1 and 2 — same conversation, English for one user, Japanese for the other.",
        "Settings: Screenshot 3 — each user chooses receive language in the side panel.",
        "Extras: voice, video, read-aloud, quality scoring, WhatsApp-style layout.",
        "Safety: API key required; Google keys stay on the server, not in the browser.",
    ])


def build() -> Document:
    doc = Document()
    setup(doc)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("UgaJapa Translation for Mattermost")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0x16, 0x6D, 0xE0)
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run("Complete Project Documentation").font.size = Pt(16)
    s2 = doc.add_paragraph()
    s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s2.add_run("Written for everyone — no coding knowledge required").italic = True
    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in ["Version 2.3.0  |  June 2026  |  UgaJapa Team"]:
        m.add_run(line)
    doc.add_page_break()

    # TOC
    heading(doc, "Table of Contents")
    for i, title in enumerate([
        "1. What is this project?",
        "2. Important clarification: what we cloned (and what we did NOT)",
        "3. The three parts of the system",
        "4. Where everything came from",
        "5. Tools and programming languages (explained simply)",
        "6. How translation works — step by step",
        "7. All features explained",
        "8. WhatsApp-style chat layout (new in version 2.2.0)",
        "9. Translation quality — how we measure accuracy",
        "10. The Translation API (company requirements)",
        "11. How to verify we are using Google Translate",
        "12. Why our translations may differ from translate.google.com",
        "13. How to run the project",
        "14. Screenshots — see the project working (with explanations)",
        "15. How to demonstrate the project (for your company)",
        "16. Questions and answers",
        "17. How to explain this project to anyone (quick script)",
        "18. Glossary — every technical word explained",
    ], 1):
        doc.add_paragraph(title, style="List Number")
    doc.add_page_break()

    # 1
    heading(doc, "1. What is this project?")
    para(doc,
         "This project adds automatic translation to Mattermost — a team chat application "
         "similar to Slack or Microsoft Teams.")
    para(doc,
         "Imagine a team where one person writes in French, another reads in English, "
         "and another in Japanese. Our system lets everyone stay in one chat channel. "
         "Each person reads messages in their own preferred language.")
    para(doc,
         "This project was built during a company internship. The team was asked to integrate "
         "real-time translation into Mattermost (TransChecker-style integration). "
         "We had to build:")
    bullets(doc, [
        "A plugin that works inside Mattermost",
        "A separate Translation API service (our own small server)",
        "Quality checking on translations",
        "Without changing Mattermost's own source code",
    ])
    note(doc, "Plugin ID:", "com.transchecker.translation")
    note(doc, "Project folder on your computer:", r"C:\Users\PC\Desktop\for mattermost")

    # 2
    heading(doc, "2. Important clarification: what we cloned (and what we did NOT)")
    note(doc, "You are NOT wrong to be confused.", 
         "Many people think we downloaded all of Mattermost. We did not.")

    heading(doc, "What we did NOT clone or download", 2)
    bullets(doc, [
        "We did NOT download Mattermost source code.",
        "There is NO folder full of Mattermost program code in our project.",
        "We did NOT modify Mattermost itself.",
    ])

    heading(doc, "What we DID use from Mattermost's website", 2)
    para(doc,
         "We used a small starter project from Mattermost's GitHub called "
         "'mattermost-plugin-starter-template'.")
    para(doc,
         "Think of it like an empty form: Mattermost gives developers a ready-made folder "
         "structure so they can build a plugin. We copied that template, then added all "
         "our translation features inside it.")
    note(doc, "Template link:", "https://github.com/mattermost/mattermost-plugin-starter-template")
    note(doc, "Our plugin folder:", "mattermost-plugin-translation/")

    heading(doc, "How Mattermost (the chat app) runs on your computer", 2)
    para(doc,
         "Mattermost runs inside Docker — a tool that downloads a ready-made application "
         "and runs it in a container (like a sealed box on your computer).")
    bullets(doc, [
        "We wrote one file: docker-compose.yml",
        "When you run docker compose up -d, Docker downloads Mattermost automatically",
        "Image used: mattermost/mattermost-team-edition version 10.5.0",
        "Mattermost opens in your browser at http://localhost:8065",
    ])
    para(doc,
         "So: Mattermost comes from Docker Hub (the internet), not from a folder we cloned.")

    # 3
    heading(doc, "3. The three parts of the system")
    tbl(doc, ["Part", "What it is (simple words)", "Where it runs"],
        [
            ["1. Mattermost", "The chat application where people send messages", "Docker, port 8065"],
            ["2. Our plugin", "Extra features added inside Mattermost (translation UI, voice, etc.)", "Inside Mattermost + your browser"],
            ["3. Translation API", "Our own small translation server we built", "Your PC, port 5000"],
        ])
    para(doc,
         "When someone sends a message, the plugin asks our Translation API to translate it. "
         "The API talks to Google Translate (or a free backup service). The result comes back "
         "and each reader sees the message in their language.")

    # 4
    heading(doc, "4. Where everything came from")
    heading(doc, "4.1 The project folder", 2)
    para(doc, "Everything lives in: for mattermost/ on your Desktop. This is YOUR team's project folder.")
    tbl(doc, ["Folder or file", "Who created it", "What it does"],
        [
            ["docker-compose.yml", "We wrote it", "Starts Mattermost + database in Docker"],
            ["translation-api/", "We created from scratch", "Our translation server"],
            ["mattermost-plugin-translation/", "We started from template, then built", "Plugin inside Mattermost"],
            ["docs/", "We wrote it", "Documentation (this file)"],
        ])

    heading(doc, "4.2 translation-api/ — our Translation API", 2)
    para(doc, "We created this folder ourselves. It is a small web server that:")
    bullets(doc, [
        "Receives text and returns a translation",
        "Detects what language text is in",
        "Gives a quality score (how good the translation probably is)",
        "Checks a password (API key) before allowing access",
        "Counts how many characters were translated (for billing demo)",
        "Also handles voice-to-text and read-aloud",
    ])
    para(doc, "To install its helper libraries, we ran npm install once. That downloaded packages into node_modules/ (normal for Node.js projects).")

    heading(doc, "4.3 Packages we installed (npm install)", 2)
    tbl(doc, ["Name", "Plain explanation"],
        [
            ["express", "Creates the web server that listens on port 5000"],
            ["dotenv", "Reads secret keys from the .env file"],
            ["multer", "Receives uploaded voice/audio files"],
            ["@xenova/transformers", "Runs artificial intelligence models on your computer"],
            ["audio-decode", "Converts voice recordings into a form the AI can understand"],
        ])

    heading(doc, "4.4 AI models (downloaded automatically the first time you use them)", 2)
    tbl(doc, ["Model name", "Size (approx.)", "Used for"],
        [
            ["Whisper (Xenova/whisper-small)", "~250 MB", "Understanding voice messages when Google is unavailable"],
            ["Embedding model (MiniLM)", "~120 MB", "Checking if two sentences mean the same thing (quality score)"],
        ])
    para(doc, "These download once into node_modules when you first transcribe audio or score quality.")

    heading(doc, "4.5 Google Cloud (we use but do not download)", 2)
    para(doc, "With a Google API key in translation-api/.env, we call Google's services over the internet:")
    bullets(doc, [
        "Google Translate — main translation",
        "Google Speech — understanding short voice messages",
        "Google Text-to-Speech — read-aloud button",
    ])
    para(doc, "If Google key is missing, text translation falls back to MyMemory (free online service).")

    heading(doc, "4.6 mattermost-plugin-translation/", 2)
    para(doc, "This plugin has two halves:")
    bullets(doc, [
        "Server half (Go language) — runs inside Mattermost, handles messages and calls our API",
        "Webapp half (TypeScript/React) — runs in the user's browser, shows buttons, translated text, voice UI, and WhatsApp-style chat colours",
    ])
    para(doc, "To install the plugin, we build a .tar.gz file and upload it in Mattermost System Console.")
    note(doc, "Latest version at time of writing:", "2.3.0 — includes WhatsApp-style layout, voice/video translation, pre-send preview, and Google Cloud integration.")

    # 5
    heading(doc, "5. Tools and programming languages (explained simply)")
    para(doc, "A programming language is how humans write instructions for computers. We used several:")
    tbl(doc, ["Language", "Used where", "In simple terms"],
        [
            ["Go (Golang)", "Plugin server folder", "Fast language for Mattermost's backend plugin — handles messages"],
            ["TypeScript", "Plugin webapp + Translation API", "Modern JavaScript with types — for websites and our API server"],
            ["JavaScript", "webapp/dist/main.js only", "Final compiled file the browser runs (auto-generated, not hand-written)"],
            ["JSON", "plugin.json, package.json", "Configuration files — lists settings in a standard format"],
            ["YAML", "docker-compose.yml", "Configuration for Docker"],
            ["Python", "docs/generate_full_documentation.py only", "Used only to create this Word document"],
        ])
    para(doc, "You do NOT need to know these languages to use or demonstrate the project.")

    heading(doc, "Other tools", 2)
    tbl(doc, ["Tool", "What it is"],
        [
            ["Docker", "Runs Mattermost in an isolated container on your PC"],
            ["Node.js + npm", "Runs our Translation API; npm installs libraries"],
            ["Google Chrome / Edge", "Best browser for voice and video recording"],
            ["Git", "Version control — tracks changes to our project folder"],
        ])

    # 6
    heading(doc, "6. How translation works — step by step")
    para(doc, "Example: Abdel writes in French. Yohana reads in English.")
    numbered(doc, [
        "Abdel types Bonjour and presses Send.",
        "Mattermost saves the original French text in the database.",
        "Our plugin notices a new message (automatic hook).",
        "Plugin checks: who is in this channel? What language does each person want?",
        "For Yohana (English), plugin sends Bonjour to our Translation API.",
        "API translates to Hello and calculates quality scores.",
        "Plugin sends Hello to Yohana's browser.",
        "Yohana sees Hello in the chat. Abdel still sees Bonjour.",
    ])
    note(doc, "Key idea:", "One original message is stored. Each person sees their own language on screen.")

    # 7
    heading(doc, "7. All features explained")
    features = [
        ("Auto-translate", "New text messages are automatically translated for each reader in their chosen language."),
        ("Receive language", "Each user picks the language they want to read messages in (channel header translate icon → side panel)."),
        ("Manual translate", "Post menu (three dots) → Translate message — forces translation for one message."),
        ("Voice messages", "Microphone button. Record up to 5 minutes. The server uses Google Speech to turn audio into text, then Google Translate for each reader. Translation appears automatically or via Translate to text."),
        ("Video messages", "Camera button. Same pipeline as voice — we extract the audio from the video, then Google Speech + Google Translate."),
        ("Read-aloud", "Speaker icon reads a message aloud. Voice can be male, female, or neutral."),
        ("Chat slang", "Short forms like bjr are expanded to bonjour before translating."),
        ("Languages panel", "Shows channel members and their languages (right side of screen)."),
        ("Slash commands", "/translation-lang fr sets your language. /translate Hello shows translation with quality scores."),
        ("Pre-send preview", "When you type in a different language than your receive setting, a popup appears before send showing how the message will be translated. Save your language in Account Settings → Translation → Save."),
        ("Caching", "Same translation is remembered for a while to avoid repeating API calls."),
        ("Usage tracking", "API counts characters translated per month (billing prototype)."),
        ("WhatsApp-style chat layout", "Messages you send appear on the right in a green bubble; messages from others appear on the left in a white bubble — like WhatsApp. See Section 8 for full details."),
    ]
    for name, desc in features:
        heading(doc, name, 2)
        para(doc, desc)

    # 8 — WhatsApp-style chat layout
    heading(doc, "8. WhatsApp-style chat layout (new in version 2.2.0)")
    para(doc,
         "In a normal Mattermost channel, every message looks the same: name on the left, "
         "avatar on the left, text below. We added an optional visual style so the chat feels "
         "more like WhatsApp — a popular phone messaging app.")
    para(doc,
         "This is only a visual change in the browser. It does NOT change how messages are saved, "
         "translated, or sent. Translation still works exactly the same underneath.")

    heading(doc, "What you see on screen", 2)
    tbl(doc, ["Who sent the message", "Where it appears", "What it looks like"],
        [
            ["You (the logged-in user)", "Right side of the chat", "Light green rounded bubble (like WhatsApp sent messages)"],
            ["Someone else in the channel", "Left side of the chat", "White rounded bubble with a thin border (like WhatsApp received messages)"],
        ])
    para(doc,
         "Your name, time, and avatar stay on the right next to your messages — the same way "
         "another person's name and avatar sit on the left next to theirs. We fixed this carefully "
         "so your name does not float in the middle of the screen.")

    heading(doc, "When you send several messages in a row", 2)
    para(doc,
         "If you send two or three messages one after another without anyone else writing in between, "
         "WhatsApp does not repeat your name and picture every time. We copied that behaviour:")
    bullets(doc, [
        "First message in a group: shows your name, time, and avatar.",
        "Next messages right after: only the green bubble — no repeated name or avatar.",
        "The bubbles still line up neatly on the right, one under the other.",
    ])

    heading(doc, "How we built it (simple explanation)", 2)
    para(doc,
         "We did NOT change Mattermost's own program files. That was a strict rule for this project. "
         "Instead, our plugin adds special styling (like a coat of paint) on top of the normal chat "
         "when Mattermost is open in your browser.")
    bullets(doc, [
        "The plugin watches the message list and marks each message as 'sent by you' or 'received from others'.",
        "Custom colours and rounded corners are applied only inside the chat area.",
        "The layout was tested and adjusted through versions 2.1.5 up to 2.2.0 until spacing, bubble width, "
        "name position, and consecutive messages all matched what we wanted.",
    ])
    note(doc, "Current plugin version:", "2.3.0 (file: dist/com.transchecker.translation-2.3.0.tar.gz)")

    heading(doc, "How to install or update this layout", 2)
    numbered(doc, [
        "Build the plugin bundle (see Section 13) or use the ready-made .tar.gz file in the dist/ folder.",
        "In Mattermost: System Console → Plugins → Plugin Management.",
        "Upload the new .tar.gz file (or disable the old plugin, upload, then enable again).",
        "In your browser, press Ctrl+Shift+R to hard refresh so the new look loads.",
        "Open a channel and send a test message — yours should appear green on the right.",
    ])

  # 9 — quality (was 8)
    heading(doc, "9. Translation quality — how we measure accuracy")
    para(doc,
         "We cannot ask a human translator every time. Instead we use the TransChecker method:")
    numbered(doc, [
        "Translate the text to the target language (e.g. French → English).",
        "Translate the result back to the original language (English → French).",
        "Compare the back-translation with what the person originally wrote.",
        "If they are similar, the translation probably preserved the meaning.",
    ])
    para(doc, "We use three checks:")
    tbl(doc, ["Score name", "What it checks"],
        [
            ["Character match (Levenshtein)", "Are the letters similar?"],
            ["Word overlap (semantic_score)", "Do the same words appear?"],
            ["AI meaning check (embedding_score)", "Do they mean the same thing even if words differ?"],
            ["quality_score", "Combined final percentage — the main number"],
        ])
    para(doc, "Where users see scores:")
    bullets(doc, [
        "Usually NOT on normal messages in the channel",
        "/translate command — shows full breakdown",
        "Pre-send preview — may show quality % for text messages",
        "Voice and video use a fast translate path (one Google call) — no quality % under voice bubbles, by design for speed",
    ])

    heading(doc, "9.1 Voice and video — how it works (same pipeline)", 2)
    para(doc,
         "Voice and video use the same steps. The only difference is that video includes a picture; "
         "we still listen to the audio track.")
    numbered(doc, [
        "User records and sends a voice or video message.",
        "Our plugin sends the audio file to the Translation API (/transcribe).",
        "Google Speech-to-Text converts speech to text (Whisper is only a backup if Google fails).",
        "Google Translate converts the text into each reader's receive language.",
        "The reader sees the translation as text under the player.",
    ])
    note(doc, "Timing:", "Voice/video often take 10–30 seconds. This is normal — speech recognition and translation are not instant.")
    note(doc, "Engines:", "Check http://localhost:5000/health — translation_engine should be google-translate and google_speech_configured should be true.")

    # 10 — API (was 9)
    heading(doc, "10. The Translation API (company requirements)")
    para(doc, "The company asked for an API that does five things. We delivered all five:")
    tbl(doc, ["Requirement", "Done?", "How"],
        [
            ["Translates text", "Yes", "POST /translate"],
            ["Detects source language", "Yes", "POST /detect"],
            ["Returns quality score", "Yes", "quality_score and related fields in response"],
            ["API-key authentication", "Yes", "Header X-API-Key must match .env"],
            ["Usage-based billing", "Yes (prototype)", "GET /usage — counts characters per month"],
        ])

    # 11 — Verify Google
    heading(doc, "11. How to verify we are using Google Translate")
    para(doc,
         "Our Translation API can use Google Translate when a Google API key is saved in "
         "translation-api/.env. If the key is missing, it falls back to a free service called MyMemory. "
         "You can check which one is active without reading any code.")
    heading(doc, "Step 1 — Make sure the Translation API is running", 2)
    para(doc, "In a terminal, inside the translation-api folder, run: npm run dev")
    heading(doc, "Step 2 — Open the health page", 2)
    para(doc, "In your browser, go to: http://localhost:5000/health")
    para(doc, "You should see something like this (the exact numbers do not matter):")
    bullets(doc, [
        '"status": "ok"',
        '"translation_engine": "google-translate"  ← this means Google is active',
        '"google_translate_configured": true  ← this means your API key is loaded',
        '"google_speech_configured": true  ← voice/video speech recognition uses Google',
    ])
    note(doc, "If you see", '"translation_engine": "mymemory" — Google key is missing or empty in .env.')
    heading(doc, "Step 3 — Test a real translation", 2)
    para(doc,
         "You can also send a test translation and look at the answer. The response includes "
         '"engine": "google-translate" when Google did the work. (This test needs the API key '
         "password from .env — ask your team member who set up the server.)")
    para(doc,
         "Important: The Mattermost plugin does not talk to Google directly. It always talks to "
         "our Translation API on port 5000 first. The API then calls Google on our behalf.")

    # 12 — Google website vs our API
    heading(doc, "12. Why our translations may differ from translate.google.com")
    para(doc,
         "It is normal if a sentence translated in our chat does not match word-for-word what "
         "you get on https://translate.google.com in your browser. That does NOT mean the system "
         "is broken. Here is why, in plain language:")
    tbl(doc, ["Reason", "Simple explanation"],
        [
            ["Different products", "The Google website and our Google Cloud API are related but not identical. Google can update the website separately."],
            ["Automatic language detection", "On the website you pick languages yourself. Our system guesses the source language automatically — short messages are easy to guess wrong."],
            ["Each user has their own read language", "You might compare your test to English, but a teammate's screen shows Japanese or French based on their settings."],
            ["Chat slang expansion", 'Our API expands shortcuts like "bjr" to "bonjour" before translating. On the website you might type the full word instead.'],
            ["Quality scoring picks the best attempt", "For tricky short messages, our API may try more than one path and keep the highest quality score — not always the same as Google's single answer."],
            ["Sender always sees the original", "The person who wrote the message still sees their own words. Only readers see the translation."],
        ])
    note(doc, "Bottom line:", 
         "Small differences are expected. To compare fairly, use the same exact text, the same target language, "
         "and confirm /health shows google-translate.")

    # 13 — How to run (was 10)
    heading(doc, "13. How to run the project")
    numbered(doc, [
        "Open terminal in translation-api folder → npm install (first time only) → npm run dev",
        "Open terminal in for mattermost folder → docker compose up -d",
        "Open browser → http://localhost:8065 → log in to Mattermost",
        "System Console → enable plugins and uploads",
        "Build plugin: cd mattermost-plugin-translation → powershell -ExecutionPolicy Bypass -File .\\scripts\\build-bundle.ps1",
        "Upload dist/com.transchecker.translation-2.3.0.tar.gz (or latest version) in Plugin Management",
        "Set Translation API URL to http://host.docker.internal:5000 and matching API key",
        "Hard refresh browser (Ctrl+Shift+R)",
    ])

    # 14 — Screenshots
    build_screenshots_section(doc)
    doc.add_page_break()

    # 15 — Demo
    heading(doc, "15. How to demonstrate the project (for your company)")
    para(doc,
         "Use Section 14 screenshots as your main visual proof. Then optionally do these live steps:")
    numbered(doc, [
        "Show Screenshots 1 and 2 together — same chat, two languages.",
        "Show Screenshot 3 — language panel and member badges (EN / JA).",
        "Show Screenshot 4 — voice and video buttons.",
        "Send a new live message from one account — show the other account sees it translated.",
        "Show the sender still sees their original text.",
        "Run /translate with a sample phrase — show quality percentages.",
        "Click a speaker icon — demonstrate read-aloud.",
        "Open http://localhost:5000/health — show translation_engine is google-translate.",
        "Explain: plugin + external API; Mattermost core was not modified.",
    ])

    # 16 — Q&A
    heading(doc, "16. Questions and answers")
    qa = [
        ("Did we clone Mattermost?", "No. We run Mattermost from Docker. We only used Mattermost's plugin starter template."),
        ("What is a plugin?", "An add-on that extends Mattermost without changing its core code."),
        ("What is an API?", "A way for programs to talk to each other over the network. Our plugin calls our Translation API."),
        ("What is Docker?", "Software that runs applications in containers. We use it to run Mattermost easily."),
        ("Where are Google keys?", "Only in translation-api/.env on the server — never in the browser."),
        ("Are we really using Google Translate?", "Yes, when GOOGLE_TRANSLATE_API_KEY is set. Check http://localhost:5000/health — it should say google-translate. See Section 11."),
        ("Why is our translation different from the Google website?", "Normal — see Section 12. Different product, auto-detection, and per-user languages all play a role."),
        ("What is the WhatsApp-style layout?", "Green bubbles on the right for your messages, white on the left for others. See Section 8."),
        ("Did we change Mattermost itself for WhatsApp layout?", "No. Only plugin styling in the browser — Mattermost core code was not touched."),
        ("Does preview always show?", "No. It only shows when you type in a language different from your receive language. You must save your language in Account Settings → Translation → Save."),
        ("Voice or video shows an error (fetch failed)?", "Usually a brief network issue with Google. Click Translate to text again. Restart translation-api if it keeps happening."),
        ("Why is voice slow?", "Google must listen to the whole clip and translate it. 10–30 seconds is normal in our setup."),
        ("Are voice and video using Google?", "Yes — Google Speech for listening, Google Translate for text. Whisper is backup only. Check /health."),
    ]
    for q, a in qa:
        heading(doc, q, 2)
        para(doc, a)

    # 17 — Explain to anyone
    build_explain_anyone_section(doc)
    doc.add_page_break()

    # 18 — Glossary
    heading(doc, "18. Glossary — every technical word explained")
    glossary = [
        ("API", "Application Programming Interface — rules for how software talks to other software."),
        ("API Key", "A secret password sent with each request to prove the caller is allowed."),
        ("Auto-translate", "Automatically showing messages in each reader's language."),
        ("Back-translation", "Translating a result back to the original language to check quality."),
        ("Bubble (chat bubble)", "The rounded coloured box around a message — green for yours, white for others in WhatsApp style."),
        ("CSS / styling", "Instructions that control colours, positions, and shapes on a web page — used for the WhatsApp look without changing Mattermost itself."),
        ("Docker", "Tool to run applications in isolated containers."),
        ("Docker Compose", "File (docker-compose.yml) that starts multiple containers together."),
        ("Docker Hub", "Website where pre-built application images are stored."),
        ("Embedding / AI score", "Computer compares meaning of two sentences using machine learning."),
        ("Go / Golang", "Programming language used for the plugin server."),
        ("Google Cloud Translation API", "Google's paid translation service our API calls — not the same as the translate.google.com website."),
        ("Hard refresh", "Ctrl+Shift+R in the browser — forces the page to reload new plugin styles."),
        ("Hook", "Automatic trigger when something happens in Mattermost (e.g. new message posted)."),
        ("KV store", "Small database inside the plugin for saving cached translations."),
        ("Mattermost", "Open-source team chat application."),
        ("MyMemory", "Free online translation service used as backup when Google key is missing."),
        ("Node.js", "Runtime that runs JavaScript/TypeScript on a server (used for Translation API)."),
        ("npm", "Tool that installs JavaScript libraries (npm install)."),
        ("Plugin", "Add-on module for Mattermost."),
        ("React", "Library for building user interfaces in the browser."),
        ("Receive language", "Language a user wants to read incoming messages in."),
        ("Redux", "Stores translation state in the browser while you use Mattermost."),
        ("STT (Speech-to-Text)", "Converting spoken audio into written text."),
        ("TransChecker", "Reference design for translation with quality checking in chat."),
        ("Translation API", "Our server on port 5000 that handles translation requests."),
        ("TTS (Text-to-Speech)", "Converting written text into spoken audio."),
        ("TypeScript", "Programming language — JavaScript with extra safety checks."),
        ("WebSocket", "Live connection that pushes translation results to the browser instantly."),
        ("WhatsApp-style layout", "Chat appearance copied from WhatsApp: your messages right/green, others left/white."),
        ("Whisper", "OpenAI speech recognition model we run locally for voice messages."),
        (".tar.gz file", "Compressed plugin package you upload in Mattermost Plugin Management."),
    ]
    for term, definition in glossary:
        p = doc.add_paragraph()
        p.add_run(f"{term}: ").bold = True
        p.add_run(definition)

    doc.add_paragraph()
    f = doc.add_paragraph()
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f.add_run("— End of Document —").italic = True

    return doc


def main() -> None:
    doc = build()
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
